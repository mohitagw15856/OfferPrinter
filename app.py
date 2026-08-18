"""OfferPrinter WebUI — a clean Streamlit app.

Run with:  streamlit run app.py

Paste or upload your CV, paste a job description or its URL, pick a provider,
and click "Print my application". Each artifact streams in with its own download
buttons. Everything runs locally with your own API key.

If no API key is configured, the app falls back to **demo mode** and shows the
bundled example package instead — which is what makes a public hosted demo
possible without asking strangers for credentials.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import streamlit as st

from offerprinter import __version__
from offerprinter.config import load_config
from offerprinter.controllers.pipeline import Pipeline
from offerprinter.models.schemas import Artifact, Locale, Provider
from offerprinter.pricing import format_cost
from offerprinter.services.cv_parser import cv_from_text, extract_cv_from_bytes
from offerprinter.services.jd_fetcher import load_job_description
from offerprinter.services.pdf_writer import markdown_to_pdf_bytes
from offerprinter.services.writer import _combined_markdown

st.set_page_config(
    page_title="OfferPrinter — free AI job application generator",
    page_icon="🖨",
    layout="centered",
    menu_items={
        "About": "OfferPrinter — one CV + one job description → a full, tailored, "
        "zero-fabrication application package. Free and open source (MIT).",
        "Report a bug": "https://github.com/mohitagw15856/OfferPrinter/issues",
    },
)

#: The bundled sample run, used for demo mode.
DEMO_DIR = Path(__file__).parent / "examples" / "output" / "northbank-senior-product-analyst"
DEMO_FILES = [
    ("Tailored CV", "tailored-cv"),
    ("Cover Letter", "cover-letter"),
    ("Fit Memo", "fit-memo"),
    ("ATS Keyword Report", "ats-keyword-report"),
    ("Interview Prep Pack", "interview-prep-pack"),
]


# --- sidebar: provider + settings ------------------------------------------


def sidebar_config():
    st.sidebar.header("⚙️ Settings")
    cfg = load_config()

    providers = [p.value for p in Provider]
    provider = st.sidebar.selectbox(
        "Provider", providers, index=providers.index(cfg.llm.provider.value)
    )
    model = st.sidebar.text_input("Model (blank = provider default)", value=cfg.llm.model)
    api_key = st.sidebar.text_input(
        "API key",
        value=cfg.llm.api_key,
        type="password",
        help="Used only to call your chosen provider. Never stored or sent anywhere else.",
    )
    locale = st.sidebar.radio(
        "English",
        [Locale.UK.value, Locale.US.value],
        index=0 if cfg.output.locale == Locale.UK else 1,
        horizontal=True,
    )

    st.sidebar.divider()
    formats = st.sidebar.multiselect(
        "Save to disk as",
        ["md", "docx", "pdf"],
        default=cfg.output.formats,
        help="Downloads below are always available regardless of this setting.",
    )
    score_fit = st.sidebar.toggle("Score the fit (0-100)", value=cfg.generation.fit_score)
    parallel = st.sidebar.toggle(
        "Generate in parallel", value=cfg.generation.parallel, help="Much faster. Same output."
    )

    st.sidebar.divider()
    st.sidebar.caption("Local-first & private: nothing leaves your machine except the LLM call.")
    st.sidebar.caption(f"OfferPrinter v{__version__} · MIT licensed")
    st.sidebar.markdown("[⭐ Star on GitHub](https://github.com/mohitagw15856/OfferPrinter)")

    cfg.llm.provider = Provider(provider)
    cfg.llm.model = model
    cfg.llm.api_key = api_key
    cfg.output.locale = Locale(locale)
    cfg.output.formats = formats or ["md"]
    cfg.generation.fit_score = score_fit
    cfg.generation.parallel = parallel
    return cfg


# --- shared rendering --------------------------------------------------------


def render_fit(fit) -> None:  # noqa: ANN001 - a FitScore
    """Show the headline score, bar, strengths and gaps."""
    colour = "green" if fit.score >= 70 else "orange" if fit.score >= 50 else "red"
    st.markdown(f"### 🎯 Fit score: :{colour}[{fit.score}/100] — {fit.band}")
    st.progress(fit.score / 100)
    st.caption(fit.verdict)
    left, right = st.columns(2)
    with left:
        if fit.strengths:
            st.markdown("**Where you genuinely match**")
            for item in fit.strengths:
                st.markdown(f"- {item}")
    with right:
        if fit.gaps:
            st.markdown("**Real gaps** (not filled in for you)")
            for item in fit.gaps:
                st.markdown(f"- {item}")


def artifact_downloads(artifact: Artifact, key_prefix: str = "") -> None:
    """Markdown / Word / PDF buttons for one artifact."""
    col1, col2, col3 = st.columns(3)
    col1.download_button(
        "⬇ Markdown",
        artifact.content,
        file_name=f"{artifact.filename}.md",
        mime="text/markdown",
        key=f"{key_prefix}md-{artifact.key}",
        use_container_width=True,
    )
    col2.download_button(
        "⬇ Word",
        _artifact_docx_bytes(artifact),
        file_name=f"{artifact.filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"{key_prefix}docx-{artifact.key}",
        use_container_width=True,
    )
    col3.download_button(
        "⬇ PDF",
        markdown_to_pdf_bytes(artifact.content),
        file_name=f"{artifact.filename}.pdf",
        mime="application/pdf",
        key=f"{key_prefix}pdf-{artifact.key}",
        use_container_width=True,
    )


# --- demo mode ----------------------------------------------------------------


def render_demo() -> None:
    """Show the bundled example package — no API key, no cost, no waiting."""
    st.info(
        "**Demo mode.** This is a real, previously generated package for an "
        "anonymised analyst applying to a *Senior Product Analyst* role. "
        "Add an API key in the sidebar to print your own.",
        icon="👀",
    )
    if not DEMO_DIR.is_dir():  # pragma: no cover - only if examples/ was removed
        st.error("The bundled example package is missing from this install.")
        return

    for title, filename in DEMO_FILES:
        path = DEMO_DIR / f"{filename}.md"
        if not path.is_file():
            continue
        artifact = Artifact(
            key=filename, title=title, filename=filename, content=path.read_text(encoding="utf-8")
        )
        with st.expander(f"📄 {title}", expanded=(filename == "fit-memo")):
            st.markdown(artifact.content)
            artifact_downloads(artifact, key_prefix="demo-")

    st.caption(
        "Notice the ATS report refuses to add 'dbt' or 'fintech' — it marks them as "
        "genuine gaps instead. That's the no-fabrication guarantee doing its job."
    )


# --- main -------------------------------------------------------------------


def main():
    st.title("🖨 OfferPrinter")
    st.markdown(
        "**Paste one job description and your CV. Get a complete, tailored "
        "application package — with zero fabrication.**"
    )

    cfg = sidebar_config()

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Your CV")
        cv_file = st.file_uploader(
            "Upload CV (.pdf, .docx, .md, .txt)", type=["pdf", "docx", "md", "txt"]
        )
        cv_text = st.text_area(
            "…or paste your CV", height=200, placeholder="Paste your CV text here"
        )
    with col2:
        st.subheader("The Job")
        jd_value = st.text_area(
            "Paste the job description, or a URL to fetch",
            height=260,
            placeholder="Paste the JD text, or https://careers.example.com/123",
        )

    action_col, roast_col = st.columns([3, 1])
    go = action_col.button("🖨  Print my application", type="primary", use_container_width=True)
    roast_me = roast_col.button("🔥 Roast my CV", use_container_width=True)

    # No key configured? Show the bundled sample instead of an error wall.
    if not cfg.llm.api_key and not (go or roast_me):
        render_demo()
        return

    if not go and not roast_me:
        return

    # --- validate inputs ---
    if not cv_file and not cv_text.strip():
        st.error("Please upload or paste your CV.")
        return
    if not cfg.llm.api_key and cfg.llm.provider is not Provider.OLLAMA:
        st.error(f"Please add an API key for {cfg.llm.provider.value} in the sidebar.")
        return

    # --- load CV ---
    try:
        if cv_file is not None:
            resume = extract_cv_from_bytes(cv_file.getvalue(), cv_file.name)
        else:
            resume = cv_from_text(cv_text)
    except Exception as exc:  # noqa: BLE001
        st.error(f"Could not read CV: {exc}")
        return

    # --- roast is CV-only, so it can run without a job description ---
    if roast_me:
        try:
            with st.spinner("Sharpening knives…"):
                artifact = Pipeline(cfg).generator.roast(resume)
        except Exception as exc:  # noqa: BLE001
            st.error(f"Roast failed: {exc}")
            return
        st.markdown(artifact.content)
        artifact_downloads(artifact, key_prefix="roast-")
        return

    if not jd_value.strip():
        st.error("Please paste a job description or a URL.")
        return

    try:
        with st.spinner("Loading job description…"):
            job = load_job_description(jd_value, timeout=cfg.llm.timeout)
    except Exception as exc:  # noqa: BLE001
        st.error(f"Could not load job description: {exc}")
        return

    # --- run pipeline, streaming artifacts ---
    pipeline = Pipeline(cfg)
    progress = st.progress(0.0, text="Starting…")
    total = max(len(cfg.generation.enabled()), 1)
    done = 0
    package = None

    try:
        for event in pipeline.stream(resume, job):
            if event.kind == "meta":
                st.success(f"🎯 Target: **{event.message}**")
            elif event.kind == "artifact":
                done += 1
                progress.progress(done / total, text=f"Generated: {event.message}")
                art = event.artifact
                with st.expander(f"📄 {art.title}", expanded=(done == 1)):
                    st.markdown(art.content)
                    artifact_downloads(art)
            elif event.kind == "fit":
                progress.progress(1.0, text="Scoring fit…")
            elif event.kind == "written":
                package = event.package
                st.info(f"💾 Also saved to: `{event.message}`")
            elif event.kind == "done":
                package = event.package
    except Exception as exc:  # noqa: BLE001
        st.error(f"Generation failed: {exc}")
        return

    progress.progress(1.0, text="Done")

    if not package:
        return

    # --- fit score ---
    if package.fit:
        st.divider()
        render_fit(package.fit)

    # --- combined downloads ---
    st.divider()
    combined_md = _combined_markdown(package)
    dcol1, dcol2, dcol3 = st.columns(3)
    dcol1.download_button(
        "⬇ Full package (Markdown)",
        combined_md,
        file_name=f"{package.slug}-full-package.md",
        mime="text/markdown",
        use_container_width=True,
    )
    dcol2.download_button(
        "⬇ Full package (PDF)",
        markdown_to_pdf_bytes(combined_md),
        file_name=f"{package.slug}-full-package.pdf",
        mime="application/pdf",
        use_container_width=True,
    )
    dcol3.download_button(
        "⬇ Everything (.zip)",
        _zip_bytes(package),
        file_name=f"{package.slug}.zip",
        mime="application/zip",
        use_container_width=True,
    )

    usage = package.usage
    if usage.calls:
        st.caption(
            f"💸 {usage.calls} calls · {usage.total_tokens:,} tokens · "
            f"{format_cost(usage.cost_usd)}"
        )
    st.caption(
        "Every line is drawn from your real CV — nothing is fabricated. Review before sending."
    )


def _artifact_docx_bytes(artifact) -> bytes:
    """Render one artifact to .docx in memory for a download button."""
    import re

    from docx import Document

    from offerprinter.services.writer import _add_markdown_runs

    doc = Document()
    for raw_line in artifact.content.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            _add_markdown_runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:])
        elif re.match(r"^\d+\.\s", line.lstrip()):
            _add_markdown_runs(
                doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", line.lstrip())
            )
        else:
            _add_markdown_runs(doc.add_paragraph(), line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _zip_bytes(package) -> bytes:
    """Bundle every artifact (Markdown + PDF) plus the combined file into a zip."""
    buf = io.BytesIO()
    combined = _combined_markdown(package)
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for art in package.artifacts:
            zf.writestr(f"{art.filename}.md", art.content.rstrip() + "\n")
            zf.writestr(f"{art.filename}.pdf", markdown_to_pdf_bytes(art.content))
        if package.fit:
            zf.writestr("fit-score.md", package.fit.as_markdown())
        zf.writestr("full-package.md", combined)
        zf.writestr("full-package.pdf", markdown_to_pdf_bytes(combined))
    return buf.getvalue()


if __name__ == "__main__":
    main()
