"""Read a CV from .pdf, .docx, .md, or plain text into clean text.

Everything downstream works on plain text, so this is the only place that knows
about file formats. All parsing is local — nothing is uploaded.
"""

from __future__ import annotations

import io
from pathlib import Path

from offerprinter.models.schemas import ResumeInput


class CVParseError(RuntimeError):
    """Raised when a CV file cannot be read."""


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs]
    # Include table cell text too — some CVs put contact details in tables.
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(part for part in parts if part.strip()).strip()


def _from_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace").strip()


_HANDLERS = {
    ".pdf": _from_pdf,
    ".docx": _from_docx,
    ".md": _from_text,
    ".markdown": _from_text,
    ".txt": _from_text,
}


def extract_cv_from_bytes(data: bytes, filename: str) -> ResumeInput:
    """Extract CV text from raw bytes, dispatching on the filename's suffix."""
    suffix = Path(filename).suffix.lower()
    handler = _HANDLERS.get(suffix)
    if handler is None:
        # Fall back to treating unknown types as UTF-8 text.
        handler = _from_text
    try:
        text = handler(data)
    except Exception as exc:  # noqa: BLE001 - surface a friendly message
        raise CVParseError(f"Could not read CV '{filename}': {exc}") from exc

    if not text.strip():
        raise CVParseError(
            f"No text could be extracted from '{filename}'. If it is a scanned PDF, "
            f"paste the CV text instead."
        )
    return ResumeInput(text=text, source=filename)


def extract_cv(path: str | Path) -> ResumeInput:
    """Read a CV from a filesystem path."""
    p = Path(path)
    if not p.is_file():
        raise CVParseError(f"CV file not found: {p}")
    return extract_cv_from_bytes(p.read_bytes(), p.name)


def cv_from_text(text: str) -> ResumeInput:
    """Wrap already-pasted CV text."""
    if not text.strip():
        raise CVParseError("The pasted CV is empty.")
    return ResumeInput(text=text.strip(), source="pasted")
